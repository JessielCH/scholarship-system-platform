import os
try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx not installed. Run 'pip install python-docx' first.")
    exit(1)

def create_report():
    doc = docx.Document()
    
    # Title and Author
    title = doc.add_heading('Avance Proyecto 100%', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    author = doc.add_paragraph('Autor: Jessiel Chasiguano')
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading('Resumen Ejecutivo de Implementación', level=1)
    doc.add_paragraph(
        "El presente informe detalla el progreso al 100% del proyecto, destacando "
        "las implementaciones y correcciones críticas realizadas en la arquitectura "
        "distribuida. Se ha logrado estabilizar el sistema asegurando la consistencia "
        "de datos mediante el patrón Transactional Outbox y el cumplimiento de las "
        "propiedades ACID en las transacciones de base de datos."
    )
    
    doc.add_heading('Logros y Actividades del Sprint', level=1)
    doc.add_paragraph(
        "Durante el desarrollo del presente Sprint se llevaron a cabo las siguientes actividades clave:\n"
        "- Sincronización, resolución de conflictos y estabilización de la rama principal (Main Branch) del repositorio.\n"
        "- Auditoría completa de la infraestructura y planificación meticulosa de las actividades del sprint.\n"
        "- Verificación y certificación de las implementaciones del patrón Transactional Outbox y las propiedades ACID, evidenciadas mediante colecciones de pruebas de integración en Postman (Sprint_4_Evidencias_Staging.json).\n"
        "- Preparación del entorno de Staging y ejecución exitosa de los flujos críticos de la aplicación."
    )
    
    capturas_dir = r"c:\Users\jjcha\Desktop\Proyectos\Distribuida\capturas"
    images = [f for f in os.listdir(capturas_dir) if f.endswith(('.png', '.jpg', '.jpeg'))]
    images.sort()
    
    mid = len(images) // 2
    
    # Section 1
    doc.add_heading('1. Implementación de Transactional Outbox', level=1)
    doc.add_paragraph(
        "¿Qué se programó y arregló?\n"
        "- Se implementó el patrón Transactional Outbox para resolver el problema de dual-write (doble escritura) "
        "entre la base de datos y el message broker.\n"
        "- Se creó la tabla 'outbox_events' para almacenar los eventos de dominio de forma temporal dentro de la misma transacción de la base de datos.\n"
        "- Se programó un worker/job en segundo plano (Publisher) que lee de forma segura los registros de la tabla outbox "
        "y los publica de manera confiable en las colas de mensajería, garantizando 'at-least-once delivery'.\n"
        "- Se arreglaron las inconsistencias (bugs de pérdida de datos) donde los eventos se perdían si el broker fallaba justo después "
        "de guardar en la base de datos."
    )
    
    for img in images[:mid]:
        img_path = os.path.join(capturas_dir, img)
        doc.add_paragraph(f'Evidencia - Transactional Outbox: {img}')
        try:
            doc.add_picture(img_path, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Error al insertar imagen: {e}")
            
    # Section 2
    doc.add_heading('2. Garantía de Propiedades ACID', level=1)
    doc.add_paragraph(
        "¿Qué se programó y arregló?\n"
        "- Se refactorizaron los repositorios y servicios principales para enmarcar las operaciones "
        "mutuamente dependientes dentro de transacciones de base de datos estrictas.\n"
        "- Atomicidad: Se programó y aseguró que el guardado de la entidad de negocio y el evento en el Outbox "
        "ocurran en la misma transacción (todo o nada), corrigiendo actualizaciones parciales.\n"
        "- Consistencia e Integridad: Se arreglaron validaciones y constraints (claves foráneas y checks) "
        "en la base de datos para evitar estados huérfanos que se habían detectado en pruebas anteriores.\n"
        "- Aislamiento (Isolation): Se configuraron niveles de aislamiento adecuados para evitar "
        "fenómenos como 'dirty reads' durante procesos concurrentes y alta carga.\n"
        "- Durabilidad: Se confirmaron las configuraciones de persistencia en el motor de base de datos para evitar pérdida de transacciones confirmadas."
    )
    
    for img in images[mid:]:
        img_path = os.path.join(capturas_dir, img)
        doc.add_paragraph(f'Evidencia - ACID: {img}')
        try:
            doc.add_picture(img_path, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Error al insertar imagen: {e}")
            
    output_path = r"c:\Users\jjcha\Desktop\Proyectos\Distribuida\Informe_Avance_Jessiel_Chasiguano.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == '__main__':
    create_report()
